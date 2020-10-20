#enconding: utf-8
from selenium import webdriver
import win32com.client as win32
from docx import Document
from docx.shared import Cm
from datetime import datetime
import win32api, sys, os
import pyautogui
import time
usuario = '99999999'
contraseña = '201803'
browser = webdriver.Chrome("C:\Program Files\chromedriver_win32\chromedriver.exe")
browser.get('https://oficinavirtual.bancoomeva.com.co/IB/presentation/bccp_mb/index.htm')
browser.set_window_size(1400,1100)
time.sleep(3)

#Coloca usuario
caja = browser.find_element_by_xpath('//*[@id="inp_555687"]')
caja.send_keys(usuario)

#Click ingreso usuario
browser.find_element_by_xpath('//*[@id="btnInicioSesion"]').click()

time.sleep(2)
#Coloca contraseña
caja = browser.find_element_by_xpath('//*[@id="passAutenticar"]')
caja.send_keys(contraseña)

browser.find_element_by_xpath('//*[@id="btnAutenticar"]').click()

time.sleep(6)

#Captura de pantalla
captura = pyautogui.screenshot()
toma = time.time( )
captura.save(str (toma)+".jpg")

time.sleep(3)

#Salir de la página
browser.find_element_by_xpath('//*[@id="imgPower"]').click()

time.sleep(3)
#word

# Creación del documento
document = Document()
# Añadimos un titulo al documento, a nivel 0
document.add_heading('Evidencia Bancoomeva', 0)

# Descripción 2 parrafo
now = datetime.now()

'''for i in [1,2]:
    p = document.add_paragraph('Evidencia  '+ str (now))
    document.add_paragraph('https://oficinavirtual.bancoomeva.com.co/IB/presentation/bccp_mb/index.htm')
    document.add_picture(str (toma)+".jpg", width=Cm(16))'''


document.save('ejemplo1.docx')
